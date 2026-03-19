"""
Example: Create a Word document with footnotes.

This example demonstrates how to use the FootnoteAdder class to create
a Word document with proper native footnotes.

Prerequisites:
    1. Run create_template.py first to generate footnote_template.docx
    2. pip install python-docx lxml
"""

from docx import Document
from footnote_adder import FootnoteAdder


def main():
    # Load the template (create it first with create_template.py)
    doc = Document("footnote_template.docx")

    # Clear template content
    doc._body.clear_content()

    # Create footnote adder
    footnote_adder = FootnoteAdder()

    # Add title
    doc.add_heading('Sample Document with Footnotes', 0)

    # Add introduction paragraph with a footnote
    p1 = doc.add_paragraph()
    p1.add_run("This document demonstrates the use of native Word footnotes ")
    p1.add_run("created programmatically with Python.")
    footnote_adder.add_footnote(p1, "", "This footnote was added using the FootnoteAdder class.")

    # Add another paragraph with multiple footnotes
    doc.add_heading('Academic Citations', level=1)

    p2 = doc.add_paragraph()
    p2.add_run("According to recent research")
    footnote_adder.add_footnote(p2, "", "Smith, J. (2024). Research Methods. Academic Press, p. 42.")
    p2.add_run(", the methodology has been widely adopted")
    footnote_adder.add_footnote(p2, "", "Johnson, M. (2023). Modern Approaches. University Press, pp. 15-18.")
    p2.add_run(" across multiple disciplines.")

    # Add a quote with citation
    doc.add_heading('Notable Quote', level=1)

    p3 = doc.add_paragraph()
    p3.add_run('"The only true wisdom is in knowing you know nothing."')
    footnote_adder.add_endnote(p3, "", "Attributed to Socrates, as recorded by Plato in the Apology.")

    # Reusable Endnote
    doc.add_heading('Reusable Endnotes Example', level=1)
    p4 = doc.add_paragraph()
    p4.add_run("First reference to the endnote.")
    footnote_adder.add_endnote(p4, "", "This is a reusable endnote text.")
    
    p5 = doc.add_paragraph()
    p5.add_run("Second reference to the exact same endnote.")
    footnote_adder.add_endnote(p5, "", "This is a reusable endnote text.")

    # Add conclusion
    doc.add_heading('Conclusion', level=1)

    p4 = doc.add_paragraph()
    p4.add_run("This example shows how easy it is to add professional footnotes ")
    p4.add_run("to Word documents generated with python-docx.")
    footnote_adder.add_footnote(p4, "", "For more information, see the project README.")

    # Save and finalize
    output_path = "example_with_footnotes.docx"
    doc.save(output_path)
    footnote_adder.finalize_footnotes(output_path)

    print(f"Document saved to {output_path}")
    print(f"Total footnotes added: {footnote_adder.footnote_id}")


if __name__ == "__main__":
    main()
