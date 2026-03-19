import os
from docx import Document
from footnote_adder import FootnoteAdder


def test_endnotes_sharing():
    """
    Test script strictly dedicated to demonstrating and validating 
    multiple capabilities of the Endnote logic (sharing, distinct texts, and multiples).
    """
    
    # Check if template exists, if not inform the user to create it
    if not os.path.exists("footnote_template.docx"):
        print("Template 'footnote_template.docx' not found. Please run 'python create_template.py' first.")
        return

    # Load the template
    doc = Document("footnote_template.docx")
    doc._body.clear_content()

    # Create footnote / endnote adder
    # You can pass endnote_style='arabic' or 'roman' here
    adder = FootnoteAdder(endnote_style='roman')

    doc.add_heading('Endnote Capabilities Testing', 0)

    # --- Scenario 1: Initializing Multiple Distinct Endnotes ---
    doc.add_heading('1. Distinct Endnotes', level=1)
    p1 = doc.add_paragraph()
    p1.add_run("This paragraph introduces the first unique endnote.")
    adder.add_endnote(p1, "", "This is the source text of Endnote #1.")
    
    p2 = doc.add_paragraph()
    p2.add_run("And here we have a totally different subject that needs its own context.")
    adder.add_endnote(p2, "", "This is the separate source text for Endnote #2.")

    # --- Scenario 2: Reusing/Sharing Endnotes across different Paragraphs ---
    doc.add_heading('2. Sharing Endnotes Accross Paragraphs', level=1)
    p3 = doc.add_paragraph()
    p3.add_run("I am making a claim that references the very first endnote again.")
    # Notice we pass the exact string as endnote 1
    adder.add_endnote(p3, "", "This is the source text of Endnote #1.")
    
    p4 = doc.add_paragraph()
    p4.add_run("Now referring back to the second endnote text from here.")
    adder.add_endnote(p4, "", "This is the separate source text for Endnote #2.")

    # --- Scenario 3: Assigning Multiple Shared Endnotes in the exact same Paragraph ---
    doc.add_heading('3. Sharing Endnotes Inside The Same Paragraph', level=1)
    p5 = doc.add_paragraph()
    p5.add_run("Here we will mention the first endnote")
    adder.add_endnote(p5, "", "This is the source text of Endnote #1.")
    p5.add_run(", proceed to introduce a new third endnote")
    adder.add_endnote(p5, "", "This is a brand new Endnote #3.")
    p5.add_run(", and immediately reference the third one again.")
    adder.add_endnote(p5, "", "This is a brand new Endnote #3.")
    
    # Save the base document 
    output_path = "test_endnotes_results.docx"
    doc.save(output_path)
    
    # Finalize by injecting the relationships and compiling the xml correctly 
    adder.finalize_footnotes(output_path)

    print(f"Endnote test document successfully saved to: {output_path}")
    print(f"Total Unique Endnotes created: {adder.endnote_id}")
    print(f"Total Unique Footnotes created: {adder.footnote_id}")


if __name__ == "__main__":
    test_endnotes_sharing()
